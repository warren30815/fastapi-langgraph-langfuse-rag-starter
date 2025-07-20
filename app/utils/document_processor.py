from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List

import pandas as pd
import PyPDF2
import tiktoken
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.config.logging_config import get_logger
from app.config.settings import settings


class DocumentProcessor:
    """Document processor for various file formats with chunking capabilities."""

    def __init__(self):
        self.logger = get_logger("document_processor")

        # Initialize tokenizer for token count calculation
        self.tokenizer = tiktoken.get_encoding("cl100k_base")

        # Initialize text splitter with default settings
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size, chunk_overlap=settings.chunk_overlap
        )

        # Supported file extensions
        self.supported_extensions = {".pdf", ".txt", ".csv", ".docx", ".md"}

    def _count_tokens(self, text: str) -> int:
        """Count tokens in text."""
        return len(self.tokenizer.encode(text))

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks using RecursiveCharacterTextSplitter."""
        chunks = self.text_splitter.split_text(text)
        return [chunk for chunk in chunks if chunk.strip()]

    async def _process_pdf(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """Process PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""

            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                text += f"\n\nPage {page_num + 1}:\n{page_text}"

            # Clean and chunk text
            text = text.strip()
            chunks = self._chunk_text(text)

            # Create document metadata
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(
                    {
                        "text": chunk,
                        "source": filename,
                        "file_type": "pdf",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "token_count": self._count_tokens(chunk),
                    }
                )

            self.logger.info(
                "PDF processed successfully",
                filename=filename,
                pages=len(pdf_reader.pages),
                chunks=len(chunks),
                total_tokens=sum(doc["token_count"] for doc in documents),
            )

            return documents

        except Exception as e:
            self.logger.error("Failed to process PDF", filename=filename, error=str(e))
            raise

    async def _process_docx(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """Process DOCX file."""
        try:
            doc = Document(BytesIO(file_content))
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Clean and chunk text
            text = text.strip()
            chunks = self._chunk_text(text)

            # Create document metadata
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(
                    {
                        "text": chunk,
                        "source": filename,
                        "file_type": "docx",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "token_count": self._count_tokens(chunk),
                    }
                )

            self.logger.info(
                "DOCX processed successfully",
                filename=filename,
                paragraphs=len(doc.paragraphs),
                chunks=len(chunks),
                total_tokens=sum(doc["token_count"] for doc in documents),
            )

            return documents

        except Exception as e:
            self.logger.error("Failed to process DOCX", filename=filename, error=str(e))
            raise

    async def _process_csv(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """Process CSV file."""
        try:
            df = pd.read_csv(BytesIO(file_content))

            # Convert DataFrame to text representation
            text = f"CSV File: {filename}\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n"
            text += f"Total Rows: {len(df)}\n\n"

            # Add column descriptions if available
            text += "Data Summary:\n"
            for col in df.columns:
                text += (
                    f"- {col}: {df[col].dtype}, non-null values: {df[col].count()}\n"
                )

            text += "\nSample Data:\n"
            text += df.head(10).to_string(index=False)

            # For numerical analysis
            if len(df.select_dtypes(include=["number"]).columns) > 0:
                text += "\n\nNumerical Summary:\n"
                text += df.describe().to_string()

            # Chunk the text
            chunks = self._chunk_text(text)

            # Create document metadata
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(
                    {
                        "text": chunk,
                        "source": filename,
                        "file_type": "csv",
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "token_count": self._count_tokens(chunk),
                        "rows": len(df),
                        "columns": list(df.columns),
                    }
                )

            self.logger.info(
                "CSV processed successfully",
                filename=filename,
                rows=len(df),
                columns=len(df.columns),
                chunks=len(chunks),
                total_tokens=sum(doc["token_count"] for doc in documents),
            )

            return documents

        except Exception as e:
            self.logger.error("Failed to process CSV", filename=filename, error=str(e))
            raise

    async def _process_text(
        self, file_content: bytes, filename: str
    ) -> List[Dict[str, Any]]:
        """Process text file (TXT, MD)."""
        try:
            # Decode text
            text = file_content.decode("utf-8")

            # Chunk the text
            chunks = self._chunk_text(text)

            # Determine file type
            file_type = Path(filename).suffix.lower()[1:]  # Remove the dot

            # Create document metadata
            documents = []
            for i, chunk in enumerate(chunks):
                documents.append(
                    {
                        "text": chunk,
                        "source": filename,
                        "file_type": file_type,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "token_count": self._count_tokens(chunk),
                    }
                )

            self.logger.info(
                "Text file processed successfully",
                filename=filename,
                file_type=file_type,
                chunks=len(chunks),
                total_tokens=sum(doc["token_count"] for doc in documents),
            )

            return documents

        except Exception as e:
            self.logger.error(
                "Failed to process text file", filename=filename, error=str(e)
            )
            raise

    async def process_file(
        self, file_content: bytes, filename: str, file_type: str = None
    ) -> List[Dict[str, Any]]:
        """Process a file and return chunked documents."""
        # Determine file type
        file_extension = Path(filename).suffix.lower()

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")

        # Route to appropriate processor
        if file_extension == ".pdf":
            return await self._process_pdf(file_content, filename)
        elif file_extension == ".docx":
            return await self._process_docx(file_content, filename)
        elif file_extension == ".csv":
            return await self._process_csv(file_content, filename)
        elif file_extension in {".txt", ".md"}:
            return await self._process_text(file_content, filename)
        else:
            raise ValueError(f"No processor available for {file_extension}")

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_extensions)


# Global document processor instance
document_processor: DocumentProcessor = DocumentProcessor()
